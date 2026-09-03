"""Parsers for pumping test field sheets (Excel template and Word sheets).

The paper layout (WiNGiN step test sheet) records readings in side by
side hourly column groups, each with Time (min), Water Level (m) and
Drawdown (m), followed by a Recovery block. Two rules from the sheets
drive the parsing:

* The recorded drawdown column is the increment between successive
  readings, not drawdown below static. Only time and water level are
  read; true drawdown is always recomputed as water level minus static
  water level.
* Discharge is often missing. The test still parses and produces water
  level and drawdown series, but a ``missing_discharge`` flag marks all
  transmissivity and yield results as pending.
* Units are read, not assumed. The sheet's own headings say what the
  numbers mean - "Time (min)", "Discharge per step (m3/h)" - and a crew
  that heads the column L/s or records the times in hours means exactly
  that. Both are converted to the canonical m3/h and minutes. A unit the
  toolkit cannot read is refused rather than guessed at: a discharge in an
  unreadable unit leaves the step pending, and a time column in one is
  dropped, because there is no pending state for time and every consumer
  would otherwise fit a curve to the wrong axis.

Times within each group are irregular (1, 2, 3 and 5 minute spacing);
nothing assumes uniform sampling.

Two recovery layouts occur on real sheets and both are handled:

* A dedicated recovery group with its own Time, Water Level and
  Recovery columns (Kuntolo sheet). Water level is read; the recovery
  increment column is ignored.
* A single shared time column with a Recovery column holding water
  levels during recovery (Dr. Timbo sheet). The recovery column is
  read against the shared times, interpreted as minutes since the pump
  stopped.
"""

from __future__ import annotations

import re
from pathlib import Path

import numpy as np

from ..models import DataFlag, PumpingStep, PumpingTest
from ..units import Quantity, convert, read_quantity, unit_from_label
from ..utils import clean_text, parse_number
from . import common

# Free-text discharge notes: "Constant Discharge of 2.93m3/h", "Q = 0.81 L/s".
# The unit is captured rather than assumed - the old pattern only matched an
# m3/h tail, so a rate written in L/s was silently dropped instead of read.
_DISCHARGE_TEXT_RE = re.compile(
    r"discharge\s*(?:of|0f)?\s*[:=]?\s*(\d+(?:\.\d+)?)\s*"
    r"([a-zµμ]{1,3}\s*3?\s*/\s*[a-z]{1,4}|lps|lpm|lph)",
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Locating the column groups
# ---------------------------------------------------------------------------

def _find_groups(grid: list[list]) -> tuple[int, list[dict]] | None:
    """Find the header row of Time / Water Level / ... column groups.

    Returns ``(row_index, groups)``. Each group has column indices and
    ``kind`` of "pumping" or "recovery". A "reco" column adjacent to a
    time column forms a recovery triplet (level read from its own water
    level column); a distant "reco" column shares the pumping time
    column and holds water levels itself.
    """
    best: tuple[int, list[dict]] | None = None
    for r, row in enumerate(grid):
        texts = common.row_text(row)
        time_cols = [c for c, t in enumerate(texts) if t.startswith("time")]
        if not time_cols:
            continue
        groups: list[dict] = []
        for gi, c in enumerate(time_cols):
            end = time_cols[gi + 1] if gi + 1 < len(time_cols) else len(texts)
            level_col = None
            reco_col = None
            for cc in range(c + 1, end):
                t = texts[cc]
                if not t:
                    continue
                if ("water" in t or t.startswith("level")) and level_col is None:
                    level_col = cc
                elif "reco" in t and reco_col is None:
                    reco_col = cc
            if level_col is None and reco_col is None:
                continue
            # The header text travels with the group so the time unit it
            # declares - "Time (min)", "Time (h)" - is read rather than assumed.
            header = texts[c]
            if reco_col is not None and level_col is not None:
                if reco_col - c <= 2:
                    # Kuntolo style triplet: Time, Water Level, Recovery increment
                    groups.append({"time": c, "level": level_col,
                                   "kind": "recovery", "time_header": header})
                else:
                    # Dr. Timbo style: shared time column; recovery column holds levels
                    groups.append({"time": c, "level": level_col,
                                   "kind": "pumping", "time_header": header})
                    groups.append({"time": c, "level": reco_col,
                                   "kind": "recovery", "time_header": header})
            elif reco_col is not None:
                groups.append({"time": c, "level": reco_col,
                               "kind": "recovery", "time_header": header})
            else:
                groups.append({"time": c, "level": level_col,
                               "kind": "pumping", "time_header": header})
        if groups and (best is None or len(groups) > len(best[1])):
            best = (r, groups)
    return best


def _read_series(
    grid: list[list], header_row: int, group: dict
) -> tuple[np.ndarray, np.ndarray, str, list[str]]:
    """Read one Time / Water Level pair, with the times put into minutes.

    Returns ``(times_min, levels, time_unit, skipped)``. ``time_unit`` is the
    unit the column declared; it is empty when nothing was declared (minutes
    assumed) and ``"?<text>"`` when the *header* declared one that could not
    be read - in which case the caller must drop the group rather than treat
    unknown units as minutes. Unlike discharge, there is no "pending" state
    for time: every consumer would silently produce wrong transmissivities.

    ``skipped`` lists individual cells whose own text could not be read; they
    cost one reading each rather than the column, and the caller reports
    them.
    """
    header = group.get("time_header", "")
    # A unit the header declares applies to the whole column, so an
    # unreadable one there costs the group. A unit written into one cell
    # applies to that cell, so an unreadable one there costs one reading -
    # otherwise a single annotated entry ("5 (approx)") threw away the test.
    declared, header_unit = unit_from_label(header, dimension="time")
    if declared and header_unit is None:
        return (np.array([], dtype=float), np.array([], dtype=float),
                f"?{declared}", [])

    times, levels, skipped = [], [], []
    unit_text = declared
    for row in grid[header_row + 1 :]:
        cell = row[group["time"]] if group["time"] < len(row) else None
        wl = parse_number(row[group["level"]]) if group["level"] < len(row) else None
        quantity = read_quantity(cell, header, dimension="time")
        if quantity.status == "absent" or wl is None:
            continue
        if quantity.status == "unknown":
            # one unreadable reading, not a broken column - but never silent
            skipped.append(str(cell))
            continue
        if quantity.unit_text:
            unit_text = quantity.unit_text
        times.append(quantity.value)
        levels.append(wl)
    return (np.array(times, dtype=float), np.array(levels, dtype=float),
            unit_text, skipped)


_STEP_LABEL_RE = re.compile(r"^\s*step\s*\d*\s*q\b", re.IGNORECASE)


def _row_unit_hint(row: list) -> str:
    """The row's own leading label, e.g. 'Discharge per step (m3/h)'."""
    for cell in row:
        text = clean_text(cell)
        if text and "discharge" in text.lower():
            return text
    return ""


def _find_step_discharges(grid: list[list]) -> dict[int, "Quantity"]:
    """Read per step discharge from 'Step n Q' labelled cells.

    Two hazards on a real sheet, both handled here:

    * The rate is written in whatever unit the crew used. The unit is taken
      from the value cell, then the label, then the row's discharge caption,
      and the value is converted to m3/h. A unit that cannot be read is
      refused, which routes the step into the existing "discharge missing,
      results pending" path rather than producing a wrong transmissivity.
    * The neighbouring cell may be the *next* step's label rather than a
      value. Scanning it blindly read "Step 2 Q (m3/h)" as a discharge of
      2 m3/h whenever the first step's box was left empty, and fitted a
      transmissivity to it. Label-shaped cells are skipped.
    """
    discharges: dict[int, Quantity] = {}
    for row in grid:
        row_hint = _row_unit_hint(row)
        for c, cell in enumerate(row):
            label = clean_text(cell)
            if not _STEP_LABEL_RE.match(label):
                continue
            num = parse_number(label.lower().split("q")[0])
            if num is None:
                continue
            for cc in range(c + 1, min(c + 3, len(row))):
                neighbour = row[cc]
                if _STEP_LABEL_RE.match(clean_text(neighbour)):
                    break  # the next step's label, not this step's value
                quantity = read_quantity(
                    neighbour, label, row_hint, dimension="flow"
                )
                if quantity.status == "absent":
                    continue
                discharges[int(num)] = quantity
                break
    return discharges


def _discharge_candidates_from_text(grid: list[list]) -> tuple[list[float], list[str]]:
    """Discharge values mentioned in free text such as
    'Constant Discharge of 2.93m3/h'.

    Returns ``(values_in_m3_per_h, unreadable_unit_texts)``. A note whose
    unit cannot be read is reported rather than converted, so it can be
    raised as a flag instead of quietly becoming a number in m3/h.
    """
    found: list[float] = []
    unreadable: list[str] = []
    for row in grid:
        for cell in row:
            if cell is None or isinstance(cell, (int, float)):
                continue
            for m in _DISCHARGE_TEXT_RE.finditer(str(cell)):
                written = m.group(2).strip()
                value = convert(float(m.group(1)), written, "m3/h", dimension="flow")
                if value is None:
                    if written not in unreadable:
                        unreadable.append(written)
                    continue
                if value not in found:
                    found.append(value)
    return found, unreadable


# Pre-printed text that names a test kind without recording which test was
# run: the template's own title mentions both, and its constant-discharge
# column labels ("Constant discharge 61-120 min") sit on every sheet whatever
# was pumped. Reading either as an answer turned step tests into constant ones.
_CONSTANT_COLUMN_LABEL_RE = re.compile(r"constant\s+discharge\s*\d")


def _sheet_test_type(grid: list[list]) -> str:
    """Look for 'STEP TEST' or 'CONSTANT DISCHARGE' banners in the sheet.

    Only decisive text counts. A cell naming both kinds is a form title
    ("PUMPING TEST FIELD SHEET (STEP / CONSTANT DISCHARGE)"), and a
    constant-discharge column label is a heading for one of the hourly
    groups - neither says which test the crew actually ran.
    """
    for row in grid:
        for cell in row:
            text = clean_text(cell).lower()
            if not text:
                continue
            step_words = "step test" in text or "step drawdown" in text
            constant_words = (
                "constant discharge" in text or "constant rate" in text
            )
            # "(STEP / CONSTANT DISCHARGE)" offers both; any mention of a step
            # alongside constant wording makes the cell a title, not an answer
            if constant_words and "step" in text:
                continue
            if step_words:
                return "step"
            if constant_words and not _CONSTANT_COLUMN_LABEL_RE.search(text):
                return "constant"
    return ""


# ---------------------------------------------------------------------------
# Assembling the PumpingTest
# ---------------------------------------------------------------------------

def _assemble(grid: list[list], source: str) -> PumpingTest:
    fields = common.extract_header_fields(grid, max_rows=len(grid))
    site = common.site_from_fields(fields, source=source)
    flags: list[DataFlag] = []

    located = _find_groups(grid)
    if located is None:
        raise ValueError(f"No Time/Water Level column groups found in {source}")
    header_row, groups = located

    def _series(kind: str) -> list[tuple[np.ndarray, np.ndarray]]:
        """Read every column group of one kind, dropping unreadable-unit ones.

        A time column whose unit cannot be read is dropped rather than taken
        as minutes: reading hours as minutes would rescale every drawdown
        curve and every transmissivity fitted to it, silently.
        """
        out = []
        for group in groups:
            if group["kind"] != kind:
                continue
            t, wl, unit_text, skipped = _read_series(grid, header_row, group)
            if unit_text.startswith("?"):
                flags.append(
                    DataFlag(
                        "error",
                        "time_unit_unknown",
                        f"The {kind} time column is headed "
                        f"'{group.get('time_header', '')}' and its unit "
                        f"'{unit_text[1:]}' could not be read, so the readings "
                        "were not used. Head the column in minutes, hours or "
                        "seconds.",
                    )
                )
                continue
            if unit_text and unit_text.lower() not in ("min", "mins", "minute",
                                                       "minutes"):
                flags.append(
                    DataFlag(
                        "info",
                        "time_unit_converted",
                        f"The {kind} times are recorded in '{unit_text}' and "
                        "have been converted to minutes for the analysis.",
                    )
                )
            if skipped:
                flags.append(
                    DataFlag(
                        "warning",
                        "time_reading_unreadable",
                        f"{len(skipped)} {kind} reading(s) carried text this "
                        "toolkit could not read as a time and were left out ("
                        + ", ".join(repr(x) for x in skipped[:3])
                        + "). Put notes outside the reading columns.",
                    )
                )
            if len(t):
                out.append((t, wl))
        return out

    pumping_series = _series("pumping")
    recovery_series = _series("recovery")

    test_type = str(fields.get("test_type", "")).strip().lower()
    stated = bool(test_type)
    if not test_type:
        test_type = _sheet_test_type(grid)
    inferred_from_shape = not test_type
    if not test_type:
        test_type = "step" if len(pumping_series) > 1 else "constant"
    if inferred_from_shape and not stated and len(pumping_series) > 1:
        flags.append(
            DataFlag(
                "info",
                "test_type_inferred",
                f"The test type cell is blank; the {len(pumping_series)} filled "
                "column groups have been read as the steps of a step test. If "
                'this was a constant discharge test, write "constant" in the '
                "test type cell so the readings are analysed as one series.",
            )
        )

    if test_type.startswith("constant") and len(pumping_series) > 1:
        # hourly column groups are one continuous series on constant tests
        t = np.concatenate([s[0] for s in pumping_series])
        wl = np.concatenate([s[1] for s in pumping_series])
        order = np.argsort(t, kind="stable")
        pumping_series = [(t[order], wl[order])]

    step_length = fields.get("step_length_min")
    discharges = _find_step_discharges(grid)

    steps: list[PumpingStep] = []
    for i, (t, wl) in enumerate(pumping_series, start=1):
        quantity = discharges.get(i)
        steps.append(
            PumpingStep(
                step_number=i,
                time_min=t,
                water_level_m=wl,
                # A refused unit leaves this None, which is the existing
                # "results pending until discharge is supplied" path - the
                # right answer, and far better than a number in the wrong unit.
                discharge_m3_per_h=quantity.value if quantity is not None else None,
                label=f"Step {i}" if len(pumping_series) > 1 else "Pumping phase",
            )
        )
        if quantity is None:
            continue
        if quantity.status == "unknown":
            flags.append(
                DataFlag(
                    "warning",
                    "discharge_unit_unknown",
                    f"Step {i} discharge is written as "
                    f"{quantity.raw_value:g} '{quantity.unit_text}', a unit "
                    "this toolkit does not recognise, so it was not used. "
                    "Record the rate in m3/h, L/s or L/min.",
                )
            )
        elif quantity.status == "converted":
            flags.append(
                DataFlag(
                    "info",
                    "discharge_unit_converted",
                    f"Step {i} discharge {quantity.raw_value:g} "
                    f"{quantity.unit_text} read as {quantity.value:.3g} m3/h.",
                )
            )
        elif quantity.status == "assumed":
            flags.append(
                DataFlag(
                    "info",
                    "discharge_unit_assumed",
                    f"Step {i} discharge {quantity.raw_value:g} carries no "
                    "unit on the sheet and was read as m3/h. Head the "
                    "discharge row with its unit to remove the assumption.",
                )
            )

    # Free-text discharge: use it only when unambiguous (one candidate, one step)
    candidates, unreadable_units = _discharge_candidates_from_text(grid)
    missing = [s for s in steps if s.discharge_m3_per_h is None]
    for written in unreadable_units:
        flags.append(
            DataFlag(
                "warning",
                "discharge_unit_unknown",
                f"A discharge note on the sheet is written in '{written}', a "
                "unit this toolkit does not recognise, so it was not used.",
            )
        )
    if candidates and missing:
        if len(candidates) == 1 and len(steps) == 1:
            steps[0].discharge_m3_per_h = candidates[0]
            flags.append(
                DataFlag(
                    "info",
                    "discharge_from_text",
                    f"Discharge {candidates[0]:g} m3/h taken from a text note on "
                    "the sheet; confirm against the measured value.",
                )
            )
        else:
            flags.append(
                DataFlag(
                    "warning",
                    "discharge_ambiguous",
                    "Discharge mentioned in sheet text ("
                    + ", ".join(f"{c:g} m3/h" for c in candidates)
                    + ") but not assigned per step; enter values in the template.",
                )
            )

    recovery_time = recovery_level = None
    if recovery_series:
        recovery_time, recovery_level = recovery_series[0]

    if recovery_time is not None:
        test_type += "+recovery"

    swl = fields.get("static_water_level_m")
    pumping_duration = None
    if steps:
        pumping_duration = float(max(s.time_min.max() for s in steps))

    test = PumpingTest(
        site=site,
        borehole_ref=str(fields.get("borehole_ref", "") or ""),
        test_type=test_type,
        static_water_level_m=swl,
        borehole_depth_m=fields.get("borehole_depth_m"),
        pump_setting_m=fields.get("pump_setting_m"),
        step_length_min=step_length,
        steps=steps,
        recovery_time_min=recovery_time,
        recovery_level_m=recovery_level,
        pumping_duration_min=pumping_duration,
        source=str(source),
    )

    # ---- data quality flags ------------------------------------------------
    if swl is None:
        flags.append(
            DataFlag(
                "error",
                "missing_static_water_level",
                "Static water level is missing; drawdown cannot be computed.",
            )
        )
    missing_q = [s.step_number for s in steps if s.discharge_m3_per_h is None]
    if missing_q:
        flags.append(
            DataFlag(
                "warning",
                "missing_discharge",
                "Discharge not recorded for step(s) "
                + ", ".join(str(n) for n in missing_q)
                + ". Drawdown and recovery curves are produced, but transmissivity "
                "and yield results are pending until discharge values are supplied.",
            )
        )
    if swl is not None:
        # The recovery limb is checked too: a recovery that overshoots the
        # static level gives negative residual drawdown, and the recovery
        # transmissivity - the one the yield prefers - is fitted through it.
        above = []
        if steps and any(np.any(s.water_level_m < swl - 0.01) for s in steps):
            above.append("pumping")
        if recovery_level is not None and np.any(recovery_level < swl - 0.01):
            above.append("recovery")
        if above:
            flags.append(
                DataFlag(
                    "warning",
                    "water_level_above_static",
                    f"Some {' and '.join(above)} water levels are above the "
                    "stated static water level, giving negative drawdown. Check "
                    "the static level and the measuring datum on the sheet.",
                )
            )
    for s in steps:
        if np.any(np.diff(s.time_min) <= 0):
            flags.append(
                DataFlag(
                    "warning",
                    "time_not_increasing",
                    f"Times are not strictly increasing in {s.label}.",
                    s.label,
                )
            )
    if test.borehole_depth_m and test.pump_setting_m:
        if test.pump_setting_m > test.borehole_depth_m:
            flags.append(
                DataFlag(
                    "warning",
                    "pump_below_borehole",
                    "Pump setting is deeper than the borehole depth.",
                )
            )
    if test.borehole_depth_m and steps:
        max_wl = max(float(np.nanmax(s.water_level_m)) for s in steps)
        if max_wl > test.borehole_depth_m:
            flags.append(
                DataFlag(
                    "warning",
                    "level_below_borehole",
                    f"Recorded water level {max_wl:.2f} m exceeds the stated "
                    f"borehole depth {test.borehole_depth_m:.0f} m; check the sheet.",
                )
            )
    test.flags = flags
    return test


# ---------------------------------------------------------------------------
# Entry points
# ---------------------------------------------------------------------------

def read_pumping_workbook(path: str | Path) -> PumpingTest:
    """Read a pumping test from the Excel template layout."""
    grid, _ = common.load_grid(path)
    return _assemble(grid, source=str(path))


def read_pumping_docx(path: str | Path) -> PumpingTest:
    """Read a pumping test from a Word field sheet (Kuntolo style).

    Paragraph text supplies the header block; the table whose header
    contains Time / Water Level column groups supplies the readings.
    """
    import docx  # python-docx

    path = Path(path)
    document = docx.Document(str(path))

    header_grid: list[list] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        parts = [p for p in re.split(r"\t+|\s{3,}", text) if p.strip()]
        header_grid.append(parts)

    best_table: list[list] | None = None
    best_count = 0
    for table in document.tables:
        grid = [[cell.text for cell in row.cells] for row in table.rows]
        located = _find_groups(grid)
        if located and len(located[1]) > best_count:
            best_table = grid
            best_count = len(located[1])
    if best_table is None:
        raise ValueError(f"No pumping test table found in {path}")

    return _assemble(header_grid + best_table, source=str(path))
