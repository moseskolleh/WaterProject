"""House-styled .docx building blocks.

Wraps python-docx with the toolkit's house style: white background,
one accent colour, clean tables, automatic figure and table numbering,
a real table of contents field, page numbers, and a reproducible save
(fixed zip timestamps so re-running on the same data gives an
identical file).

House language rules are enforced here: em and en dashes are replaced
with plain hyphens and a lint helper flags contractions so template
text stays in plain professional English.
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

import docx
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..config import HouseStyle

_CONTRACTION_RE = re.compile(
    r"\b(?:can't|won't|don't|doesn't|didn't|isn't|aren't|wasn't|weren't|"
    r"hasn't|haven't|hadn't|shouldn't|couldn't|wouldn't|it's|that's|there's|"
    r"what's|let's|they're|we're|you're|i'm|he's|she's|who's|it'll|we'll|"
    r"you'll|they'll|i've|we've|you've|they've)\b",
    re.IGNORECASE,
)


def lint_text(text: str) -> list[str]:
    """Return house style violations in a piece of report text."""
    problems = []
    if "—" in text or "–" in text:
        problems.append("em or en dash present")
    for match in _CONTRACTION_RE.finditer(text):
        problems.append(f"contraction '{match.group(0)}'")
    return problems


def _clean(text: str) -> str:
    """Apply the house language rules to outgoing text."""
    return text.replace("—", " - ").replace("–", "-")


def _hex_to_rgb(color: str) -> RGBColor:
    color = color.lstrip("#")
    return RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))


class ReportBuilder:
    """Assemble a styled report document section by section."""

    def __init__(self, style: HouseStyle | None = None, title: str = "Report"):
        self.style = style or HouseStyle()
        self.doc = docx.Document()
        self.title = title
        self._figure_no = 0
        self._table_no = 0
        self._setup_styles()
        self._setup_page()

    # ------------------------------------------------------------------
    # set-up
    # ------------------------------------------------------------------

    def _setup_styles(self) -> None:
        accent = _hex_to_rgb(self.style.accent_color)
        normal = self.doc.styles["Normal"]
        normal.font.name = self.style.font_name
        normal.font.size = Pt(self.style.base_font_size_pt)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

        for level, size, bold in ((1, 14, True), (2, 12, True), (3, 11, True)):
            h = self.doc.styles[f"Heading {level}"]
            h.font.name = self.style.font_name
            h.font.size = Pt(size)
            h.font.bold = bold
            h.font.color.rgb = accent
            h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True

        caption = self.doc.styles["Caption"] if "Caption" in [s.name for s in self.doc.styles] else None
        if caption is not None:
            caption.font.name = self.style.font_name
            caption.font.size = Pt(9)
            caption.font.bold = True
            caption.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def _setup_page(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        for side in ("top", "bottom", "left", "right"):
            setattr(section, f"{side}_margin", Cm(2.5))
        self._add_page_numbers(section)

    def _add_page_numbers(self, section) -> None:
        footer_para = section.footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run()
        for element, text in (("begin", None), (None, "PAGE"), ("end", None)):
            if element:
                fld = OxmlElement("w:fldChar")
                fld.set(qn("w:fldCharType"), element)
                run._r.append(fld)
            else:
                instr = OxmlElement("w:instrText")
                instr.set(qn("xml:space"), "preserve")
                instr.text = " PAGE "
                run._r.append(instr)
        run.font.size = Pt(9)

    # ------------------------------------------------------------------
    # content
    # ------------------------------------------------------------------

    def cover(
        self,
        title_lines: list[str],
        subtitle_lines: list[str] | None = None,
        details: list[tuple[str, str]] | None = None,
    ) -> None:
        """Simple clean cover: organisation, title block, detail lines."""
        if self.style.logo_path and Path(self.style.logo_path).exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(self.style.logo_path, width=Cm(4.5))
        if self.style.organisation:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_clean(self.style.organisation))
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = _hex_to_rgb(self.style.accent_color)
            if self.style.organisation_details:
                p2 = self.doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.add_run(_clean(self.style.organisation_details))
                r2.font.size = Pt(9)
        for _ in range(4):
            self.doc.add_paragraph()
        for line in title_lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_clean(line))
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = _hex_to_rgb(self.style.accent_color)
        for line in subtitle_lines or []:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_clean(line))
            run.font.size = Pt(13)
        if details:
            for _ in range(3):
                self.doc.add_paragraph()
            for label, value in details:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"{_clean(label)}: ")
                run.font.bold = True
                p.add_run(_clean(str(value)))
        self.page_break()

    def table_of_contents(self) -> None:
        self.heading("Table of Contents", level=1, numbered=False)
        para = self.doc.add_paragraph()
        run = para.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = (
            "Right-click and choose Update Field to fill the table of contents."
        )
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
            run._r.append(el)
        self.page_break()

    def executive_summary(
        self, paragraphs: list[str], key_findings: list[str] | None = None
    ) -> None:
        """A short verdict block placed at the very top of a report.

        Composed from figures the report already computes, so a ministry
        engineer or programme manager gets the yield, water safety and the
        single next action before any detail. ``key_findings`` renders as a
        tight bullet list under a bold label.
        """
        self.heading("Executive Summary", level=1, numbered=False)
        for text in paragraphs:
            if not text:
                continue
            self.paragraph(text, align="justify")
        if key_findings:
            self.paragraph("Key findings:", bold=True)
            self.bullets([k for k in key_findings if k])
        self.page_break()

    def heading(self, text: str, level: int = 1, numbered: bool = True) -> None:
        self.doc.add_heading(_clean(text), level=level)

    def paragraph(self, text: str, bold: bool = False, italic: bool = False,
                  align: str | None = None):
        p = self.doc.add_paragraph()
        run = p.add_run(_clean(text))
        run.font.bold = bold
        run.font.italic = italic
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "justify":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def bullets(self, items: list[str]) -> None:
        for item in items:
            self.doc.add_paragraph(_clean(item), style="List Bullet")

    def figure(self, image_path: str | Path, caption: str, width_cm: float = 15.0) -> int:
        """Insert an image with an automatic 'Figure N.' caption.

        Returns the figure number so body text can reference it.
        """
        self._figure_no += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Cm(width_cm))
        cap = self.doc.add_paragraph()
        run = cap.add_run(f"Figure {self._figure_no}. {_clean(caption)}")
        run.font.size = Pt(9)
        run.font.bold = True
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return self._figure_no

    @property
    def next_figure_number(self) -> int:
        return self._figure_no + 1

    @property
    def next_table_number(self) -> int:
        return self._table_no + 1

    def table(
        self,
        rows: list[list],
        header: list[str] | None = None,
        caption: str | None = None,
        col_widths_cm: list[float] | None = None,
        font_size_pt: float = 9.5,
    ) -> int:
        """Insert a styled table; returns the table number.

        The caption goes above the table (report convention). Cell
        values may contain newlines for stacked layer entries.
        """
        self._table_no += 1
        if caption:
            cap = self.doc.add_paragraph()
            run = cap.add_run(f"Table {self._table_no}. {_clean(caption)}")
            run.font.size = Pt(9)
            run.font.bold = True
        n_cols = len(header) if header else max(len(r) for r in rows)
        table = self.doc.add_table(rows=0, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        accent = self.style.accent_color.lstrip("#")

        if header:
            cells = table.add_row().cells
            for i, text in enumerate(header):
                cells[i].text = ""
                para = cells[i].paragraphs[0]
                run = para.add_run(_clean(str(text)))
                run.font.bold = True
                run.font.size = Pt(font_size_pt)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:fill"), accent)
                cells[i]._tc.get_or_add_tcPr().append(shading)
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                if i >= n_cols:
                    break
                cells[i].text = ""
                para = cells[i].paragraphs[0]
                run = para.add_run(_clean("" if value is None else str(value)))
                run.font.size = Pt(font_size_pt)
        if col_widths_cm:
            for row in table.rows:
                for i, width in enumerate(col_widths_cm[:n_cols]):
                    row.cells[i].width = Cm(width)
        self.doc.add_paragraph()
        return self._table_no

    def header_block_table(self, pairs: list[tuple[str, str]], font_size_pt: float = 9.5) -> None:
        """Two label/value pairs per row, like the field sheet headers."""
        table = self.doc.add_table(rows=0, cols=4)
        table.style = "Table Grid"
        for i in range(0, len(pairs), 2):
            row = table.add_row().cells
            for j, (label, value) in enumerate(pairs[i : i + 2]):
                cell_label = row[j * 2]
                cell_value = row[j * 2 + 1]
                cell_label.text = ""
                run = cell_label.paragraphs[0].add_run(_clean(label))
                run.font.bold = True
                run.font.size = Pt(font_size_pt)
                cell_value.text = ""
                run_v = cell_value.paragraphs[0].add_run(_clean(str(value)))
                run_v.font.size = Pt(font_size_pt)
        self.doc.add_paragraph()

    def signature_block(
        self, name: str, role: str, phone: str = "", organisation: str = "", email: str = ""
    ) -> None:
        self.paragraph("REPORT SUBMITTED BY:", bold=True)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.paragraph("." * 30)
        self.paragraph(name, bold=True)
        p = self.paragraph(role)
        p.runs[0].font.italic = True
        if organisation:
            self.paragraph(organisation)
        if phone:
            self.paragraph(f"Cell: {phone}")
        if email:
            self.paragraph(f"Email: {email}")

    def page_break(self) -> None:
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------------------------
    # save
    # ------------------------------------------------------------------

    def save(self, path: str | Path) -> Path:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        core = self.doc.core_properties
        core.title = self.title
        core.author = self.style.organisation or "Groundwater Toolkit"
        self.doc.save(path)
        _normalise_zip_timestamps(path)
        return path


def _normalise_zip_timestamps(path: Path) -> None:
    """Rewrite the .docx zip with fixed timestamps and sorted entries so
    identical content produces an identical file."""
    with zipfile.ZipFile(path, "r") as zin:
        entries = sorted(zin.namelist())
        contents = {name: zin.read(name) for name in entries}
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in entries:
            info = zipfile.ZipInfo(name, date_time=(2020, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            zout.writestr(info, contents[name])
