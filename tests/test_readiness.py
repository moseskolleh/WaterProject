"""The certification gate.

A completion report is what a borehole is handed over on. The toolkit will
write one from whatever it has - an interim report is a real need - but a
document resting on missing evidence must not be indistinguishable from one
that is not. These tests pin the two halves of that: what counts as
established, and what the document says when it is not.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from groundwater.design import design_borehole
from groundwater.hydraulics import analyse_pumping_test
from groundwater.ingestion import (
    read_drilling_workbook,
    read_pumping_workbook,
    read_quality_workbook,
)
from groundwater.models import SiteMetadata, WaterQualityResult
from groundwater.quality import assess_sample
from groundwater.readiness import REPORTS, REQUIREMENTS, assess_readiness


@pytest.fixture(scope="module")
def project(sample_data) -> dict:
    """A complete Dr Timbo project, with a position the sheets do not carry."""
    d = Path(sample_data) / "dr_timbo"
    log = read_drilling_workbook(d / "dr_timbo_drilling_log.xlsx")
    analysis = analyse_pumping_test(read_pumping_workbook(d / "dr_timbo_constant_test.xlsx"))
    sample = read_quality_workbook(d / "dr_timbo_water_quality.xlsx")
    return {
        "site": SiteMetadata(
            community="Dr. Timbo's", district="Western Area Rural",
            easting=778000.0, northing=946000.0, utm_zone=28,
        ),
        "drilling_log": log,
        "pump_analysis": analysis,
        "wq_assessment": assess_sample(sample),
        "borehole_design": design_borehole(
            log=log, static_water_level_m=analysis.test.static_water_level_m),
    }


def test_a_complete_project_is_ready(project):
    readiness = assess_readiness(project, "completion")
    assert readiness.state == "ready"
    assert readiness.is_certifiable
    assert readiness.unmet == []


def test_an_empty_project_is_not_ready():
    """Fail closed: nothing loaded is not the same as nothing wrong.

    Two of the checks are negative ones - no unreadable unit, no fatal flag -
    and they are vacuously true when nothing has been loaded. That is
    harmless: every substantive requirement beside them is unmet, so the
    project is not certifiable and the report says exactly which pieces of
    evidence are missing.
    """
    readiness = assess_readiness({}, "completion")
    assert readiness.state == "not_ready"
    assert not readiness.is_certifiable
    outstanding = {r.key for r in readiness.unmet}
    assert outstanding == {
        "site_located", "borehole_logged", "pumping_measured",
        "yield_established", "water_quality_panel", "water_quality_evaluable",
        "design_derived",
    }


def test_the_gate_judges_evidence_not_the_answer(project, sample_data):
    """A borehole that fails a health guideline is still certifiable.

    The finding is the point of the report. A gate that held back bad news
    would teach people to leave the bad news out.
    """
    assessment = project["wq_assessment"]
    assert assessment.verdict_state == "health_fail"     # the sample really does
    assert assess_readiness(project, "completion").state == "ready"


def test_an_ungraded_determinand_holds_the_report_back(project, sample_data):
    """Even when the verdict is already a failure.

    A health exceedance outranks uncertainty in the verdict, and rightly, but
    it does not make an unreadable result readable - so the row is checked,
    not the headline.
    """
    d = Path(sample_data) / "dr_timbo"
    sample = read_quality_workbook(d / "dr_timbo_water_quality.xlsx")
    sample.results.append(WaterQualityResult("Iron", 0.1, "wibbles"))
    state = dict(project, wq_assessment=assess_sample(sample))

    readiness = assess_readiness(state, "completion")
    assert readiness.state == "not_ready"
    outstanding = {r.key for r in readiness.unmet}
    assert "water_quality_evaluable" in outstanding
    assert "Iron" in dict((r.key, r.detail) for r in readiness.unmet)[
        "water_quality_evaluable"]


def test_a_missing_position_holds_every_report_back(project):
    state = dict(project, site=SiteMetadata(community="Nowhere"))
    # the sheets carry no coordinates either, so nothing can supply one
    for report in ("completion", "handover", "quality", "pumping"):
        readiness = assess_readiness(state, report)
        assert "site_located" in {r.key for r in readiness.unmet}, report


def test_a_position_on_any_sheet_locates_the_project(project):
    """Crews write the GPS on one sheet and not the others."""
    log = project["drilling_log"]
    log.site.easting, log.site.northing, log.site.utm_zone = 778000.0, 946000.0, 28
    try:
        state = dict(project, site=SiteMetadata(community="Dr. Timbo's"))
        readiness = assess_readiness(state, "completion")
        assert "site_located" not in {r.key for r in readiness.unmet}
    finally:
        log.site.easting = log.site.northing = None


def test_a_pumping_report_does_not_need_the_water_quality(project):
    """Each report is judged on what it actually claims."""
    state = dict(project, wq_assessment=None)
    assert assess_readiness(state, "pumping").state == "ready"
    assert assess_readiness(state, "completion").state == "not_ready"


def test_an_override_issues_the_report_without_certifying_it(project):
    state = dict(project, wq_assessment=None)
    readiness = assess_readiness(
        state, "completion",
        overrides={
            "water_quality_panel": {"reason": "interim; lab result awaited",
                                    "by": "M. Kolleh"},
            "water_quality_evaluable": "interim; lab result awaited",
        },
    )
    assert readiness.state == "ready_with_overrides"
    # issuable, but never certifiable - those are different claims
    assert not readiness.is_certifiable
    assert {r.key for r in readiness.overridden} == {
        "water_quality_panel", "water_quality_evaluable"}
    assert "M. Kolleh" in [r.override_by for r in readiness.overridden]
    assert "Issued on override" in readiness.summary


def test_an_override_cannot_invent_a_requirement_that_was_met(project):
    """Overriding something already established must not change its record."""
    readiness = assess_readiness(
        project, "completion", overrides={"borehole_logged": "not needed"})
    logged = [r for r in readiness.requirements if r.key == "borehole_logged"][0]
    assert logged.state == "met" and not logged.override_reason


def test_a_check_that_raises_is_not_a_pass():
    """A broken check must never read as an established requirement."""
    class Exploding:
        flags = property(lambda self: (_ for _ in ()).throw(RuntimeError("boom")))

    readiness = assess_readiness({"pump_analysis": Exploding()}, "pumping")
    assert readiness.state == "not_ready"
    assert any("could not run" in r.detail for r in readiness.unmet)


def test_the_readiness_record_is_serialisable(project):
    import json

    import yaml

    payload = assess_readiness(project, "completion").as_dict()
    assert json.dumps(payload)
    assert yaml.safe_dump(payload)


def test_every_report_names_only_requirements_that_exist():
    for report, keys in REPORTS.items():
        assert keys, report
        for key in keys:
            assert key in REQUIREMENTS, f"{report} names an unknown requirement {key}"


def test_a_provisional_report_says_so_on_its_cover(project, tmp_path):
    """The stamp is the whole point: it travels with the document."""
    from docx import Document

    from groundwater.reporting.quality import QualityReportInputs, build_quality_report

    state = dict(project, site=SiteMetadata(community="Nowhere"))
    readiness = assess_readiness(state, "quality")
    assert not readiness.is_certifiable

    out = build_quality_report(
        QualityReportInputs(
            assessment=project["wq_assessment"], figures_dir=tmp_path,
            include_diagrams=False, readiness=readiness,
        ),
        tmp_path / "quality.docx",
    )
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "PROVISIONAL - NOT FOR CERTIFICATION" in text
    assert "Site position" in text


def test_a_ready_report_carries_no_stamp(project, tmp_path):
    from docx import Document

    from groundwater.reporting.quality import QualityReportInputs, build_quality_report

    readiness = assess_readiness(project, "quality")
    assert readiness.is_certifiable
    out = build_quality_report(
        QualityReportInputs(
            assessment=project["wq_assessment"], figures_dir=tmp_path,
            include_diagrams=False, readiness=readiness,
        ),
        tmp_path / "quality_ready.docx",
    )
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "PROVISIONAL" not in text
    assert "NOT A CERTIFICATION" not in text


def test_an_overridden_report_names_who_issued_it(project, tmp_path):
    from docx import Document

    from groundwater.reporting.quality import QualityReportInputs, build_quality_report

    state = dict(project, site=SiteMetadata(community="Nowhere"))
    readiness = assess_readiness(
        state, "quality",
        overrides={"site_located": {"reason": "GPS unit failed; position to "
                                              "follow", "by": "M. Kolleh"}},
    )
    assert readiness.state == "ready_with_overrides"
    out = build_quality_report(
        QualityReportInputs(
            assessment=project["wq_assessment"], figures_dir=tmp_path,
            include_diagrams=False, readiness=readiness,
        ),
        tmp_path / "quality_override.docx",
    )
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "ISSUED ON OVERRIDE" in text
    assert "M. Kolleh" in text
    assert "GPS unit failed" in text
