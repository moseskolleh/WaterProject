import hashlib
import re

import docx
import pytest

from groundwater.config import Config
from groundwater.design import design_borehole
from groundwater.hydraulics import analyse_pumping_test
from groundwater.ingestion import (
    check_all,
    read_drilling_workbook,
    read_pumping_workbook,
    read_quality_workbook,
    read_ves_workbook,
)
from groundwater.quality import assess_sample
from groundwater.reporting import build_geophysical_report
from groundwater.reporting.completion import CompletionReportInputs, build_completion_report
from groundwater.reporting.docx_utils import lint_text
from groundwater.reporting.geophysical import GeophysicalReportInputs
from groundwater.reporting.handover import HandoverReportInputs, build_handover_report
from groundwater.reporting.pumping import PumpingReportInputs, build_pumping_report
from groundwater.reporting.quality import QualityReportInputs, build_quality_report
from groundwater.ves import interpret_model, invert_sounding


def _document_text(path) -> str:
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


@pytest.fixture(scope="module")
def geophysical_report(sample_data, tmp_path_factory):
    tmp = tmp_path_factory.mktemp("geo")
    soundings = read_ves_workbook(sample_data / "rokel" / "rokel_ves.xlsx")
    inversions = [invert_sounding(s) for s in soundings]
    interps = [interpret_model(s, r.model) for s, r in zip(soundings, inversions)]
    inputs = GeophysicalReportInputs(
        soundings=soundings,
        inversions=inversions,
        interpretations=interps,
        figures_dir=tmp,
        geologist_name="Test Geologist",
        flags=check_all([(s.sounding_id, s.site) for s in soundings]),
        include_qa_annex=True,
    )
    path = build_geophysical_report(inputs, tmp / "geo.docx")
    return path, inputs, tmp


def test_geophysical_report_structure(geophysical_report):
    path, _, _ = geophysical_report
    text = _document_text(path)
    for expected in (
        "Table of Contents",
        "1. Introduction",
        "2. Background and Geology of the Project Area",
        "3.1 Reconnaissance Survey",
        "3.2.1 Resistivity Profiling",
        "3.2.2 Selection of VES Points",
        "3.2.3 Vertical Electrical Sounding (VES)",
        "4. Data Analysis and Interpretation",
        "order of preference for drilling",
        "5. Conclusions and Recommendations",
        "REPORT SUBMITTED BY:",
    ):
        assert expected in text, f"missing section: {expected}"
    # every figure caption is numbered and referenced in the body
    figures = re.findall(r"Figure (\d+)\.", text)
    assert figures and figures == sorted(figures, key=int)
    assert "Figure 1" in text


def test_geophysical_report_reproducible(geophysical_report):
    path, inputs, tmp = geophysical_report
    first = hashlib.sha256(path.read_bytes()).hexdigest()
    again = build_geophysical_report(inputs, tmp / "geo2.docx")
    assert hashlib.sha256(again.read_bytes()).hexdigest() == first


def test_house_style_no_dashes_or_contractions(geophysical_report):
    path, _, _ = geophysical_report
    text = _document_text(path)
    assert "—" not in text  # em dash
    assert "–" not in text  # en dash
    assert not lint_text(text), lint_text(text)


def test_all_other_reports_build(sample_data, tmp_path):
    log = read_drilling_workbook(sample_data / "dr_timbo" / "dr_timbo_drilling_log.xlsx")
    test = read_pumping_workbook(sample_data / "dr_timbo" / "dr_timbo_constant_test.xlsx")
    sample = read_quality_workbook(sample_data / "dr_timbo" / "dr_timbo_water_quality.xlsx")
    analysis = analyse_pumping_test(test)
    assessment = assess_sample(sample)
    design = design_borehole(log=log, static_water_level_m=test.static_water_level_m)

    pumping = build_pumping_report(
        PumpingReportInputs(analysis=analysis, figures_dir=tmp_path),
        tmp_path / "pumping.docx",
    )
    completion = build_completion_report(
        CompletionReportInputs(
            log=log, design=design, pumping=analysis, quality=assessment,
            figures_dir=tmp_path,
        ),
        tmp_path / "completion.docx",
    )
    handover = build_handover_report(
        HandoverReportInputs(
            site=log.site, log=log, design=design, pumping=analysis,
            quality=assessment, figures_dir=tmp_path,
        ),
        tmp_path / "handover.docx",
    )
    quality = build_quality_report(
        QualityReportInputs(assessment=assessment, figures_dir=tmp_path),
        tmp_path / "quality.docx",
    )
    for path, must_contain in (
        (pumping, ["Cooper-Jacob", "Limitations and Uncertainty",
                   "References", "Glossary and Abbreviations"]),
        (completion, ["Borehole Log Data", "References", "Glossary and Abbreviations"]),
        (handover, ["Operation and Maintenance", "References",
                    "Glossary and Abbreviations"]),
        (quality, ["Ionic Balance", "Limitations and Uncertainty",
                   "References", "Glossary and Abbreviations"]),
    ):
        text = _document_text(path)
        for token in must_contain:
            assert token in text, f"{token} missing from {path.name}"
        assert "—" not in text
        assert not lint_text(text)
    # a real citation appears, not just a named method
    assert "Cooper, H.H. and Jacob" in _document_text(pumping)
    # traceability: the safety factor is stated explicitly
    assert "safety factor" in _document_text(pumping).lower()


def test_geophysical_report_has_limitations_references_glossary(geophysical_report):
    text = _document_text(geophysical_report[0])
    for token in ("6. Limitations and Uncertainty", "equivalence and suppression",
                  "References", "Glossary and Abbreviations", "geoBoundaries"):
        assert token in text, f"{token} missing"


def test_pending_pumping_report(sample_data, tmp_path):
    test = read_pumping_workbook(sample_data / "kuntolo" / "kuntolo_step_test.xlsx")
    analysis = analyse_pumping_test(test)
    path = build_pumping_report(
        PumpingReportInputs(analysis=analysis, figures_dir=tmp_path),
        tmp_path / "kuntolo.docx",
    )
    text = _document_text(path)
    assert "pending" in text.lower()
    assert "discharge" in text.lower()


def test_extraction_review_workbook(tmp_path):
    from groundwater.extraction import (
        ExtractedDocument,
        ExtractedField,
        ExtractedTable,
        UncertainCell,
        fill_ves_template,
        write_review_workbook,
    )
    from openpyxl import load_workbook

    document = ExtractedDocument(
        source="scan.jpg",
        document_kind="ves",
        header=[
            ExtractedField("Community", "Rokel", 0.97),
            ExtractedField("GPS Coordinate East", "0708958", 0.6),
        ],
        tables=[
            ExtractedTable(
                title="Schlumberger Array VES Field Data",
                columns=["No.", "AB/2 (m)", "MN (m)", "Apparent Resistivity (ohm-m)"],
                rows=[["1", "1", "0.4", "1165"], ["2", "2", "0.4", "1193"]],
            )
        ],
        uncertain_cells=[UncertainCell(0, 1, 3, "smudged digits")],
        extractor="claude",
    )
    review = write_review_workbook(document, tmp_path / "review.xlsx")
    wb = load_workbook(review)
    assert set(wb.sheetnames) == {"Header", "Table 1", "Review"}
    review_texts = [row[0].value for row in wb["Review"].iter_rows() if row[0].value]
    assert any("GPS Coordinate East" in t for t in review_texts)

    template = fill_ves_template(document, tmp_path / "ves.xlsx")
    filled = load_workbook(template)
    ws = filled.active
    assert ws["D2"].value == "Rokel"
    assert ws["D4"].value == "0708958"
    assert ws.cell(row=12, column=4).value == "1193"
    # then the standard parser reads the filled template
    from groundwater.ingestion import read_ves_workbook

    soundings = read_ves_workbook(template)
    assert len(soundings) == 1 and soundings[0].n_readings == 2


def test_fill_ves_template_accepts_already_canonical_field_names(tmp_path):
    """The text-layer extractor resolves labels before it stores them.

    Its header fields are therefore named by the canonical key, and a
    canonical key does not match its own label pattern ("sounding_id" is not
    "sounding number"). Feeding those names back through match_label alone
    dropped every one of them on the floor, silently.
    """
    from openpyxl import load_workbook

    from groundwater.extraction import (
        ExtractedDocument,
        ExtractedField,
        ExtractedTable,
        fill_ves_template,
    )

    document = ExtractedDocument(
        source="sheet.pdf",
        document_kind="ves",
        header=[
            ExtractedField("community", "Rokel", 0.95),
            ExtractedField("sounding_id", "VES A-1", 0.95),
            ExtractedField("client", "Living Water International", 0.95),
        ],
        tables=[
            ExtractedTable(
                title="Table 1",
                columns=["No.", "AB/2 (m)", "MN (m)", "Apparent Resistivity (ohm-m)"],
                rows=[["1", "1.5", "0.5", "210.4"], ["2", "2", "0.5", "233.1"]],
            )
        ],
        extractor="pdf-text",
    )
    template = fill_ves_template(document, tmp_path / "ves.xlsx")
    ws = load_workbook(template).active
    assert ws["D2"].value == "Rokel"
    assert ws["D3"].value == "VES A-1"
    assert ws["B2"].value == "Living Water International"

    from groundwater.ingestion import read_ves_workbook

    sounding = read_ves_workbook(template)[0]
    assert sounding.sounding_id == "VES A-1"
    assert sounding.site.community == "Rokel"


def test_pdf_kind_guess():
    from groundwater.extraction.pdf_text import _guess_kind

    assert _guess_kind("Schlumberger array VES data, apparent resistivity") == "ves"
    assert _guess_kind("constant discharge pumping test, drawdown") == "pumping_test"
    assert _guess_kind("nothing recognisable") == "unknown"


def _typed_field_sheet(path):
    """A VES field sheet as a real PDF: Helvetica text placed with Tm/Tj.

    Written here rather than committed as a binary so the fixture says out
    loud what it is - and, crucially, what it is *not*: there is not one
    ruling line on the page. The columns are held by whitespace, exactly as
    they are on a sheet typed in a word processor, which is the case
    ``extract_tables()`` cannot see.
    """
    import zlib

    rows = [[1, 1.5, 0.5, 210.4], [2, 2, 0.5, 233.1], [3, 3, 0.5, 268.0],
            [4, 4, 0.5, 291.7], [5, 6, 0.5, 302.5], [6, 8, 0.5, 288.2],
            [7, 10, 0.5, 264.9], [8, 15, 0.5, 198.3]]
    placed = [
        (60, 760, "SCHLUMBERGER ARRAY VES FIELD DATA"),
        (60, 735, "Community: Rokel"),
        (300, 735, "Client: Living Water International"),
        (60, 715, "District: Port Loko"),
        (300, 715, "Sounding Number: VES A-1"),
        (60, 695, "Date: 2023-05-14"),
        (300, 695, "Field Supervisor: M. Kolleh"),
        (60, 650, "No."), (110, 650, "AB/2 (m)"), (200, 650, "MN (m)"),
        (280, 650, "Apparent Resistivity (ohm-m)"),
    ]
    for i, row in enumerate(rows):
        y = 630 - i * 18
        for x, value in zip((60, 110, 200, 280), row):
            placed.append((x, y, str(value)))

    ops = ["BT", "/F1 10 Tf"]
    for x, y, text in placed:
        escaped = str(text).replace("\\", "\\\\").replace("(", r"\(").replace(")", r"\)")
        ops.append(f"1 0 0 1 {x} {y} Tm ({escaped}) Tj")
    ops.append("ET")
    content = zlib.compress("\n".join(ops).encode("latin-1"))

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length 6 0 R /Filter /FlateDecode >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        str(len(content)).encode(),
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"
    xref = len(out)
    out += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode()
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref}\n%%EOF\n").encode()
    path.write_bytes(bytes(out))
    return path


def test_pdf_extraction_reads_an_unruled_field_sheet(tmp_path):
    """The common case: a typed sheet with no ruling lines around its table.

    ``extract_text()`` collapses the gap between the two header columns, so a
    reader working from it gets "Rokel Client: Living Water International" as
    the community; ``extract_tables()`` needs ruling lines, so it returns
    nothing at all. Both are read from word positions instead.
    """
    pytest.importorskip("pdfplumber")
    from groundwater.extraction.pdf_text import extract_pdf_text

    document = extract_pdf_text(_typed_field_sheet(tmp_path / "ves_sheet.pdf"))

    assert document.document_kind == "ves"
    fields = {f.name: f.value for f in document.header}
    # each side of the page is its own field, not one glued to the next
    assert fields["community"] == "Rokel"
    assert fields["client"] == "Living Water International"
    assert fields["district"] == "Port Loko"
    assert fields["sounding_id"] == "VES A-1"
    assert fields["supervisor"] == "M. Kolleh"

    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.columns == ["No.", "AB/2 (m)", "MN (m)", "Apparent Resistivity (ohm-m)"]
    assert len(table.rows) == 8
    assert table.rows[0] == ["1", "1.5", "0.5", "210.4"]
    assert table.rows[-1] == ["8", "15", "0.5", "198.3"]
    # every cell parsed, so nothing is held back for review
    assert document.uncertain_cells == []
    # and the header block above it was not swept into the reading table
    assert not any("Community" in cell for row in table.rows for cell in row)


def test_pdf_extraction_feeds_the_ves_template(tmp_path):
    """End to end: an unruled PDF becomes a workbook the normal reader takes."""
    pytest.importorskip("pdfplumber")
    from groundwater.extraction import fill_ves_template
    from groundwater.extraction.pdf_text import extract_pdf_text
    from groundwater.ingestion import read_ves_workbook

    document = extract_pdf_text(_typed_field_sheet(tmp_path / "ves_sheet.pdf"))
    template = fill_ves_template(document, tmp_path / "ves.xlsx")
    sounding = read_ves_workbook(template)[0]

    assert sounding.sounding_id == "VES A-1"
    assert sounding.site.community == "Rokel"
    assert len(sounding.ab2) == 8
    assert sounding.ab2[0] == pytest.approx(1.5)
    assert sounding.rho_app[0] == pytest.approx(210.4)
    assert sounding.ab2[-1] == pytest.approx(15.0)
    assert sounding.rho_app[-1] == pytest.approx(198.3)
