"""Planned against actual, and what that makes payable.

Three ways a drilling contract loses money, and one test group each: work
measured that nobody authorised, work paid for twice, and retention nobody
tracked. In every case the requirement is the same - the money does not move
until somebody has signed for it, and where it cannot move the certificate
says so on its face rather than netting the problem away.
"""

from __future__ import annotations

import pytest

from groundwater.procurement import (
    Contract,
    ContractLine,
    Measurement,
    Variation,
    certify,
    contract_from_estimate,
    contract_summary,
    value_work,
)


def _contract(**terms) -> Contract:
    return Contract(
        ref="WSD/2024/017", contractor="WiNGiN", client="District Council",
        date="2024-02-01",
        lines=[
            ContractLine("MOB", "Mobilisation", "sum", 1, 3000.0),
            ContractLine("DRL-OB", "Drilling, overburden", "m", 20, 45.0),
            ContractLine("DRL-RK", "Drilling, rock", "m", 25, 80.0),
            ContractLine("CAS", "uPVC casing", "m", 45, 22.0),
        ],
        **terms,
    )


def test_the_contract_sum_is_the_priced_bill():
    contract = _contract()
    assert contract.sum_usd == 3000 + 20 * 45 + 25 * 80 + 45 * 22


def test_a_contract_is_frozen_from_the_estimate_not_linked_to_it(sample_data):
    """The estimate keeps moving as the design changes; a contract does not."""
    from groundwater.costing.model import CostingInputs, estimate_borehole_cost

    estimate = estimate_borehole_cost(CostingInputs(total_depth_m=45.0))
    contract = contract_from_estimate(estimate, ref="WSD/2024/017")
    assert len(contract.lines) == len(estimate.items)
    # frozen at the contract price the same estimate recommends, not at the
    # contractor's bare cost rates (which paid a third below that price)
    assert contract.sum_usd == pytest.approx(estimate.price_usd)
    assert contract.rate_basis == "price" and contract.rate_uplift > 1.0
    at_cost = contract_from_estimate(estimate, ref="X", basis="cost")
    assert at_cost.sum_usd == pytest.approx(estimate.direct_cost_usd)
    with_vat = contract_from_estimate(estimate, ref="X", basis="price_with_vat")
    assert with_vat.sum_usd == pytest.approx(estimate.price_with_vat_usd)
    with pytest.raises(ValueError):
        contract_from_estimate(estimate, ref="X", basis="retail")

    # a later change to the estimate cannot move what was signed
    estimate.items[0].quantity *= 2
    assert contract.sum_usd != pytest.approx(estimate.price_usd)


# ---------------------------------------------------------------------------
# Work nobody authorised
# ---------------------------------------------------------------------------

def test_drilling_past_the_contract_depth_is_not_payable_because_it_happened():
    """The classic. The water was deeper; somebody still has to sign for it."""
    lines, problems = value_work(
        _contract(), [Measurement("DRL-RK", 42.0)])
    rock = [line for line in lines if line.code == "DRL-RK"][0]
    assert rock.measured_quantity == 42.0
    assert rock.authorised_quantity == 25.0
    assert rock.payable_quantity == 25.0
    assert rock.overmeasure_quantity == 17.0
    assert rock.overmeasure_amount_usd == 17.0 * 80.0
    assert any("not payable until a variation authorises it" in p
               for p in problems)


def test_a_variation_makes_the_same_work_payable():
    lines, problems = value_work(
        _contract(), [Measurement("DRL-RK", 42.0)],
        [Variation("VO-1", "2024-03-04", "DRL-RK", quantity_delta=17.0,
                   reason="water struck at 62 m, not 45 m",
                   authorised_by="M. Kolleh")])
    rock = [line for line in lines if line.code == "DRL-RK"][0]
    assert rock.authorised_quantity == 42.0
    assert rock.payable_quantity == 42.0
    assert rock.overmeasure_quantity == 0.0
    assert rock.variation_refs == ("VO-1",)
    assert problems == []


def test_an_unsigned_variation_is_a_request_not_an_instruction():
    _, problems = value_work(
        _contract(), [], [Variation("VO-2", "2024-03-04", "DRL-RK",
                                    quantity_delta=17.0)])
    assert any("names nobody who authorised it" in p for p in problems)


def test_an_unsigned_variation_does_not_authorise_the_work_it_asks_for():
    """A warning beside the total is read once; the certificate is banked."""
    lines, problems = value_work(
        _contract(), [Measurement("DRL-RK", 42.0)],
        [Variation("VO-2", "2024-03-04", "DRL-RK", quantity_delta=17.0,
                   reason="water struck at 62 m, not 45 m")])
    rock = [line for line in lines if line.code == "DRL-RK"][0]
    assert rock.authorised_quantity == 25.0      # the contract, and no more
    assert rock.payable_quantity == 25.0
    assert rock.overmeasure_quantity == 17.0
    assert rock.variation_refs == ()             # nothing signed to cite
    assert any("left out of this valuation" in p for p in problems)


def test_an_unsigned_variation_does_not_reprice_a_line_either():
    lines, _ = value_work(
        _contract(), [Measurement("CAS", 45.0)],
        [Variation("VO-6", "2024-03-04", "CAS", rate_usd=26.0,
                   reason="steel casing substituted")])
    casing = [line for line in lines if line.code == "CAS"][0]
    assert casing.rate_usd == 22.0
    assert casing.payable_amount_usd == 45 * 22.0


def test_work_added_only_by_an_unsigned_variation_is_authorised_by_nobody():
    lines, problems = value_work(
        _contract(), [Measurement("GRAVEL", 12.0)],
        [Variation("VO-7", "2024-03-04", "GRAVEL", quantity_delta=12.0,
                   rate_usd=15.0, item="Gravel pack", unit="m3")])
    added = [line for line in lines if line.code == "GRAVEL"][0]
    assert added.item == "Gravel pack"           # still named, so it is visible
    assert added.authorised_quantity == 0.0
    assert added.payable_amount_usd == 0.0
    assert added.overmeasure_quantity == 12.0
    assert any("nobody has signed it" in p for p in problems)


def test_work_measured_against_nothing_at_all_is_valued_at_zero():
    lines, problems = value_work(_contract(), [Measurement("GRAVEL", 12.0)])
    stray = [line for line in lines if line.code == "GRAVEL"][0]
    assert stray.in_contract is False
    assert stray.payable_amount_usd == 0.0
    assert any("neither in the contract nor in any variation" in p
               for p in problems)


def test_a_variation_that_adds_work_without_a_rate_is_valued_at_zero():
    lines, problems = value_work(
        _contract(), [Measurement("GRAVEL", 12.0)],
        [Variation("VO-3", "2024-03-04", "GRAVEL", quantity_delta=12.0,
                   authorised_by="M. Kolleh", item="Gravel pack")])
    added = [line for line in lines if line.code == "GRAVEL"][0]
    assert added.authorised_quantity == 12.0
    assert added.payable_amount_usd == 0.0
    assert any("without giving a rate" in p for p in problems)


def test_a_variation_can_change_a_rate_and_can_omit_work():
    lines, _ = value_work(
        _contract(), [Measurement("CAS", 45.0), Measurement("MOB", 1.0)],
        [Variation("VO-4", "2024-03-04", "CAS", rate_usd=26.0,
                   reason="steel casing substituted", authorised_by="M. K."),
         Variation("VO-5", "2024-03-04", "MOB", quantity_delta=-1.0,
                   reason="rig already on site", authorised_by="M. K.")])
    casing = [line for line in lines if line.code == "CAS"][0]
    mobilisation = [line for line in lines if line.code == "MOB"][0]
    assert casing.rate_usd == 26.0 and casing.payable_amount_usd == 45 * 26.0
    assert mobilisation.authorised_quantity == 0.0
    assert mobilisation.payable_amount_usd == 0.0


def test_a_negative_measurement_is_refused_rather_than_credited():
    lines, problems = value_work(_contract(), [Measurement("CAS", -10.0)])
    casing = [line for line in lines if line.code == "CAS"][0]
    assert casing.measured_quantity == 0.0
    assert any("negative quantity" in p for p in problems)


def test_a_repeated_measurement_is_the_latest_figure_not_a_running_total():
    """They are cumulative to date; adding them up pays for the same metres twice."""
    lines, _ = value_work(
        _contract(),
        [Measurement("CAS", 20.0, "2024-03-01"),
         Measurement("CAS", 45.0, "2024-04-01")])
    casing = [line for line in lines if line.code == "CAS"][0]
    assert casing.measured_quantity == 45.0


# ---------------------------------------------------------------------------
# Paying twice
# ---------------------------------------------------------------------------

def test_a_certificate_pays_only_what_the_last_one_did_not():
    contract = _contract(retention_percent=0.0)
    first = certify(contract, [Measurement("MOB", 1.0)], number=1,
                    date="2024-03-01")
    assert first.due_now_usd == 3000.0

    second = certify(
        contract, [Measurement("MOB", 1.0), Measurement("DRL-OB", 20.0)],
        number=2, date="2024-04-01",
        previously_certified_usd=first.net_certified_usd)
    assert second.gross_usd == 3000.0 + 20 * 45
    assert second.due_now_usd == 20 * 45


def test_a_later_certificate_with_nothing_previously_certified_is_challenged():
    """Silence here pays for the first certificate's work a second time."""
    certificate = certify(_contract(), [Measurement("MOB", 1.0)], number=3,
                          date="2024-05-01")
    assert any("will be paid for that work twice" in p
               for p in certificate.problems)


def test_a_certificate_that_would_pay_less_than_nothing_pays_nothing():
    """A negative payment is a credit note somebody has to agree, not a payment."""
    contract = _contract(retention_percent=0.0)
    certificate = certify(contract, [Measurement("MOB", 1.0)], number=2,
                          date="2024-04-01",
                          previously_certified_usd=5000.0)
    assert certificate.net_certified_usd == 3000.0
    assert certificate.due_now_usd == 0.0
    assert certificate.overpaid_usd == 2000.0
    assert "nothing is due" in certificate.summary


def test_a_negative_previously_certified_is_refused():
    certificate = certify(_contract(), [], number=2, date="2024-04-01",
                          previously_certified_usd=-100.0)
    assert certificate.previously_certified_usd == 0.0
    assert any("cannot be right" in p for p in certificate.problems)


# ---------------------------------------------------------------------------
# Retention and the advance
# ---------------------------------------------------------------------------

def test_retention_is_withheld_and_capped():
    contract = _contract(retention_percent=10.0, retention_cap_percent=5.0)
    everything = [Measurement("MOB", 1.0), Measurement("DRL-OB", 20.0),
                  Measurement("DRL-RK", 25.0), Measurement("CAS", 45.0)]
    early = certify(contract, [Measurement("MOB", 1.0)], number=1,
                    date="2024-03-01")
    assert early.retention_usd == pytest.approx(300.0)      # 10% of 3000

    full = certify(contract, everything, number=2, date="2024-06-01")
    assert full.gross_usd == contract.sum_usd
    # 10% of the work would be more than the 5% cap on the contract sum
    assert full.retention_usd == pytest.approx(contract.sum_usd * 0.05)


def test_the_advance_is_recovered_in_proportion_to_the_work_done():
    contract = _contract(advance_percent=20.0, retention_percent=0.0)
    half = contract.sum_usd / 2
    certificate = certify(
        contract, [Measurement("MOB", 1.0), Measurement("DRL-OB", 20.0),
                   Measurement("DRL-RK", 25.0), Measurement("CAS", 45.0)],
        number=1, date="2024-06-01")
    assert certificate.advance_recovered_usd == pytest.approx(
        contract.advance_usd)          # all the work done, all of it recovered

    part = certify(contract, [Measurement("MOB", 1.0)], number=1,
                   date="2024-03-01")
    assert part.advance_recovered_usd == pytest.approx(
        contract.advance_usd * 3000.0 / contract.sum_usd)
    assert part.advance_recovered_usd < contract.advance_usd
    assert half > 0


def test_no_advance_means_nothing_to_recover():
    certificate = certify(_contract(), [Measurement("MOB", 1.0)], number=1,
                          date="2024-03-01")
    assert certificate.advance_recovered_usd == 0.0


# ---------------------------------------------------------------------------
# What the certificate says
# ---------------------------------------------------------------------------

def test_the_revised_sum_is_the_contract_plus_its_variations():
    contract = _contract()
    certificate = certify(
        contract, [], number=1, date="2024-04-01",
        variations=[Variation("VO-1", "2024-03-04", "DRL-RK",
                              quantity_delta=17.0, authorised_by="M. K.")])
    assert certificate.variation_usd == pytest.approx(17.0 * 80.0)
    assert certificate.revised_sum_usd == pytest.approx(
        contract.sum_usd + 17.0 * 80.0)


def test_progress_is_measured_against_the_revised_sum_not_the_original():
    contract = _contract(retention_percent=0.0)
    certificate = certify(
        contract, [Measurement("MOB", 1.0), Measurement("DRL-OB", 20.0),
                   Measurement("DRL-RK", 42.0), Measurement("CAS", 45.0)],
        number=1, date="2024-06-01",
        variations=[Variation("VO-1", "2024-03-04", "DRL-RK",
                              quantity_delta=17.0, authorised_by="M. K.")])
    assert certificate.percent_complete == pytest.approx(100.0)
    assert certificate.gross_usd == pytest.approx(certificate.revised_sum_usd)


def test_the_summary_line_carries_the_number_and_the_money():
    certificate = certify(_contract(), [Measurement("MOB", 1.0)], number=1,
                          date="2024-03-01")
    assert "Certificate 1" in certificate.summary
    assert "$2,700" in certificate.summary      # 3000 less 10% retention


def test_the_head_of_the_certificate_shows_every_deduction():
    contract = _contract(advance_percent=20.0)
    certificate = certify(contract, [Measurement("MOB", 1.0)], number=2,
                          date="2024-04-01", previously_certified_usd=500.0)
    rows = dict(contract_summary(contract, certificate))
    assert rows["Contract"] == "WSD/2024/017"
    assert rows["Contractor"] == "WiNGiN"
    assert "Less retention (10%)" in rows
    assert "Less advance recovery (20% advance)" in rows
    assert rows["Less previously certified"] == "-$500"
    assert "Due on this certificate" in rows


def test_the_certificate_record_is_serialisable():
    import json

    import yaml

    payload = certify(
        _contract(), [Measurement("DRL-RK", 42.0)], number=1, date="2024-04-01",
    ).as_dict()
    assert json.dumps(payload)
    assert yaml.safe_dump(payload)
    assert payload["problems"]
    assert len(payload["lines"]) == 4


# ---------------------------------------------------------------------------
# The document
# ---------------------------------------------------------------------------

def test_the_certificate_shows_every_deduction_on_its_face(tmp_path):
    """One that shows only the bottom figure is one nobody can dispute."""
    from docx import Document

    from groundwater.reporting.procurement import (
        PaymentCertificateInputs,
        build_payment_certificate,
    )

    contract = _contract(advance_percent=20.0)
    certificate = certify(
        contract, [Measurement("MOB", 1.0), Measurement("DRL-RK", 42.0)],
        number=2, date="2024-04-01", previously_certified_usd=1500.0,
        variations=[Variation("VO-1", "2024-03-04", "DRL-RK",
                              quantity_delta=10.0, reason="deeper water",
                              authorised_by="M. Kolleh")])
    out = build_payment_certificate(
        PaymentCertificateInputs(contract=contract, certificate=certificate,
                                 prepared_by="M. Kolleh"),
        tmp_path / "ipc2.docx")
    document = Document(str(out))
    text = "\n".join(p.text for p in document.paragraphs)
    body = "\n".join(cell.text for t in document.tables for row in t.rows
                     for cell in row.cells)

    assert "Interim Payment Certificate No. 2" in text
    assert "Before the figures" in text            # the problems come first
    assert "not payable until a variation authorises it" in text
    assert "Measured beyond what was authorised" in text
    assert "usually was" in text                   # and why it is held back
    assert "VO-1" in body                          # the variation is itemised
    for label in ("Contract sum", "Less retention (10%)",
                  "Less advance recovery (20% advance)",
                  "Less previously certified", "Due on this certificate"):
        assert label in body, label


def test_a_clean_certificate_carries_no_problem_section(tmp_path):
    from docx import Document

    from groundwater.reporting.procurement import (
        PaymentCertificateInputs,
        build_payment_certificate,
    )

    contract = _contract()
    certificate = certify(contract, [Measurement("MOB", 1.0)], number=1,
                          date="2024-03-01")
    out = build_payment_certificate(
        PaymentCertificateInputs(contract=contract, certificate=certificate),
        tmp_path / "ipc1.docx")
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Before the figures" not in text
    assert "Measured beyond what was authorised" not in text
    assert "Certification" in text


def test_a_rate_only_variation_is_a_variation():
    """Repricing a line used to leave 'Variations +$0' on the certificate
    while paying the new rate: the contract amount was valued at the varied
    rate too, so the difference cancelled."""
    contract = Contract(ref="C", lines=[
        ContractLine(code="MOB", item="Mobilisation", unit="sum", quantity=1, rate_usd=3000),
        ContractLine(code="CAS", item="Casing", unit="m", quantity=45, rate_usd=22),
    ])
    repriced = Variation(ref="VO-1", date="2024-05-01", code="CAS", rate_usd=26,
                         reason="supplier price", authorised_by="Engineer")
    cert = certify(contract, [Measurement("MOB", 1), Measurement("CAS", 45)],
                   number=1, date="2024-06-01", variations=[repriced])
    assert cert.variation_usd == pytest.approx(45 * (26 - 22))
    assert cert.revised_sum_usd == pytest.approx(3000 + 45 * 26)
    assert cert.gross_usd == pytest.approx(cert.revised_sum_usd)
    assert cert.percent_complete == pytest.approx(100.0)


def test_an_omission_beyond_the_line_authorises_nothing_and_says_so():
    contract = Contract(ref="C", lines=[
        ContractLine(code="CAS", item="Casing", unit="m", quantity=45, rate_usd=22),
    ])
    omission = Variation(ref="VO-2", date="2024-05-01", code="CAS", quantity_delta=-60,
                         reason="redesign", authorised_by="Engineer")
    lines, problems = value_work(contract, [Measurement("CAS", 10)], [omission])
    line = lines[0]
    assert line.authorised_quantity == 0.0
    assert line.payable_quantity == 0.0 and line.overmeasure_quantity == 10.0
    assert line.percent_complete is None or line.percent_complete >= 0
    assert any("omits 60" in p and "only carries 45" in p for p in problems)


def test_duplicate_contract_codes_are_named_and_valued_once():
    contract = Contract(ref="C", lines=[
        ContractLine(code="CAS", item="Casing 6 in", unit="m", quantity=10, rate_usd=22),
        ContractLine(code="CAS", item="Casing 4 in", unit="m", quantity=10, rate_usd=22),
    ])
    lines, problems = value_work(contract, [Measurement("CAS", 10)])
    assert len(lines) == 1 and lines[0].payable_amount_usd == pytest.approx(220)
    assert any("more than one line coded CAS" in p for p in problems)
