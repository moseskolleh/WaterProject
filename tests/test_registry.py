"""The asset registry.

Two things are being pinned here. The identifier has to be wrong loudly
rather than quietly, because a mistyped one silently attaches a repair
history to a different village's borehole. And the state derived from the
event stream has to fail closed: an asset register earns its reputation by
going green everywhere as soon as people stop filing reports.
"""

from __future__ import annotations

import csv
from datetime import date
from pathlib import Path

import pytest

from groundwater import registry
from groundwater.models import SiteMetadata
from groundwater.registry import (
    Asset,
    AssetEvent,
    Schedule,
    asset_state,
    check_character,
    merge_events,
    mint_asset_id,
    parse_asset_id,
    registry_stats,
    validate_asset_id,
)

# Dr. Timbo's Residence, Western Area Rural: 8.48310 N, 13.22940 W.
TIMBO = SiteMetadata(community="Dr. Timbo's", district="Western Area Rural",
                     easting=694912.0, northing=938150.0, utm_zone=28)


# ---------------------------------------------------------------------------
# The identifier
# ---------------------------------------------------------------------------

def test_the_same_borehole_gets_the_same_identifier_from_either_phone():
    """There is no server to allocate from, so it has to be derivable."""
    other = SiteMetadata(community="different name entirely",
                         district="Western Area Rural",
                         easting=694914.0, northing=938147.0, utm_zone=28)
    assert mint_asset_id(TIMBO) == mint_asset_id(other)


def test_two_boreholes_a_village_apart_get_different_identifiers():
    far = SiteMetadata(district="Western Area Rural", easting=695512.0,
                       northing=938150.0, utm_zone=28)
    assert mint_asset_id(TIMBO) != mint_asset_id(far)


def test_the_identifier_says_where_it_is_before_you_decode_it():
    assert mint_asset_id(TIMBO).startswith("SL-WAR-")


def test_a_district_nobody_recognises_is_marked_not_guessed():
    odd = SiteMetadata(district="Freetown City", easting=694912.0,
                       northing=938150.0, utm_zone=28)
    assert mint_asset_id(odd).startswith("SL-XX-")


def test_a_borehole_with_no_position_cannot_be_given_an_identifier():
    with pytest.raises(ValueError, match="nothing to find it by"):
        mint_asset_id(SiteMetadata(community="Nowhere"))


def test_every_district_in_the_data_has_a_code_and_no_two_share_one():
    """The codes are written out, so nothing checks them but this."""
    path = Path(registry.__file__).parent / "data" / "sl_districts.csv"
    with path.open(encoding="utf-8") as handle:
        districts = [row["district"] for row in csv.DictReader(handle)]
    for district in districts:
        assert district.lower() in registry.DISTRICT_CODES, district
    codes = list(registry.DISTRICT_CODES.values())
    assert len(set(codes)) == len(codes), "two districts share a code"
    assert registry.UNKNOWN_DISTRICT not in codes


def test_every_single_mistyped_character_is_caught():
    """The point of the check character. All of them, not most of them."""
    identifier = mint_asset_id(TIMBO)
    body = identifier.replace("-", "")
    alphabet = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    missed = []
    for position in range(len(body)):
        for char in alphabet:
            if char == body[position]:
                continue
            broken = body[:position] + char + body[position + 1:]
            candidate = f"SL-{broken[2:5]}-{broken[5:12]}-{broken[12:]}"
            if parse_asset_id(candidate) == identifier:
                continue  # a forgiven Crockford confusion, not a wrong code
            if parse_asset_id(candidate) is not None:
                missed.append(candidate)
    assert missed == [], missed


def test_every_transposition_of_two_adjacent_characters_is_caught():
    identifier = mint_asset_id(TIMBO)
    body = identifier.replace("-", "")
    for i in range(len(body) - 1):
        if body[i] == body[i + 1]:
            continue
        swapped = body[:i] + body[i + 1] + body[i] + body[i + 2:]
        candidate = f"SL-{swapped[2:5]}-{swapped[5:12]}-{swapped[12:]}"
        assert parse_asset_id(candidate) is None or \
            parse_asset_id(candidate) == identifier, candidate


def test_the_reading_confusions_are_forgiven_but_wrong_codes_are_not():
    """I and L for 1, O for 0, U for V - read off a plate, not different codes."""
    identifier = mint_asset_id(TIMBO)
    typed = identifier.replace("1", "I").replace("0", "O").lower()
    assert parse_asset_id(typed) == identifier
    assert parse_asset_id(" " + identifier.replace("SL-", "") + " ") == identifier
    assert parse_asset_id("SL-WAR-8T4KQ2A-7") is None or \
        parse_asset_id("SL-WAR-8T4KQ2A-7").startswith("SL-WAR-")


def test_a_wrong_identifier_is_refused_with_something_a_person_can_act_on():
    identifier = mint_asset_id(TIMBO)
    wrong = identifier[:-1] + ("Z" if identifier[-1] != "Z" else "Y")
    ok, reason = validate_asset_id(wrong)
    assert not ok
    assert "check character" in reason and "mistyped" in reason
    assert identifier[-1] in reason        # it says what it should have been

    ok, reason = validate_asset_id("")
    assert not ok and "No identifier" in reason
    ok, reason = validate_asset_id("just some words")
    assert not ok and "SL-WAR" in reason


def test_the_check_character_ignores_separators():
    assert check_character("SL-WAR-8T4KQ2A") == check_character("SLWAR8T4KQ2A")


# ---------------------------------------------------------------------------
# The event stream
# ---------------------------------------------------------------------------

def _asset(*events) -> Asset:
    return Asset(asset_id=mint_asset_id(TIMBO), community="Dr. Timbo's",
                 district="Western Area Rural", easting=694912.0,
                 northing=938150.0, utm_zone=28, events=list(events))


def test_a_borehole_nobody_has_reported_on_is_not_working_it_is_unknown():
    """The failure mode of every asset register: silence read as success."""
    state = asset_state(_asset(), date(2024, 6, 1))
    assert state.function == "unknown"
    assert not state.is_working
    assert state.label == "Not known"


def test_a_failure_with_no_repair_stays_a_failure():
    state = asset_state(_asset(
        AssetEvent("2020-01-10", "commissioned"),
        AssetEvent("2023-04-02", "failure", "rising main parted"),
    ), date(2024, 6, 1))
    assert state.function == "non_functional"
    assert state.days_out_of_service == 426
    assert not state.is_working


def test_a_repair_puts_it_back_in_service():
    state = asset_state(_asset(
        AssetEvent("2020-01-10", "commissioned"),
        AssetEvent("2023-04-02", "failure", "rising main parted"),
        AssetEvent("2023-05-11", "restored", "new rising main", by="A. Bangura"),
    ), date(2024, 6, 1))
    assert state.function == "functional"
    assert state.days_out_of_service is None
    assert "A. Bangura" in state.detail


def test_the_order_events_arrive_in_does_not_change_the_answer():
    """They arrive from two phones, in whatever order the merge happens."""
    events = [
        AssetEvent("2023-05-11", "restored"),
        AssetEvent("2020-01-10", "commissioned"),
        AssetEvent("2023-04-02", "failure"),
    ]
    forwards = asset_state(_asset(*events), date(2024, 6, 1))
    backwards = asset_state(_asset(*reversed(events)), date(2024, 6, 1))
    assert forwards.as_dict() == backwards.as_dict()


def test_a_record_dated_in_the_future_settles_nothing_yet():
    state = asset_state(_asset(
        AssetEvent("2020-01-10", "commissioned"),
        AssetEvent("2023-04-02", "failure"),
        AssetEvent("2099-01-01", "restored"),
    ), date(2024, 6, 1))
    assert state.function == "non_functional"


def test_an_unreadable_date_is_kept_but_establishes_nothing():
    """Losing a maintenance record to a typo is worse than not counting it."""
    asset = _asset(AssetEvent("last rains", "repair", "changed the seals"))
    state = asset_state(asset, date(2024, 6, 1))
    assert state.function == "unknown"
    assert state.undated_events == 1
    assert "nobody can read" in state.detail
    assert len(merge_events(asset.asset_id, asset.events)) == 1


def test_the_same_visit_written_down_twice_merges_into_one():
    """Two people at one wellhead, two phones, one repair."""
    mine = [AssetEvent("2023-05-11", "repair", "new seals", by="A. Bangura")]
    theirs = [AssetEvent("2023-05-11", "repair", "new seals", by="A. Bangura")]
    merged = merge_events("SL-WAR-8T4KQ2A-7", mine, theirs)
    assert len(merged) == 1


def test_merging_is_a_union_that_does_not_care_about_order_or_repetition():
    a = [AssetEvent("2023-05-11", "repair", "new seals")]
    b = [AssetEvent("2023-04-02", "failure", "rising main parted")]
    once = [e.event_id for e in merge_events("X", a, b)]
    again = [e.event_id for e in merge_events("X", b, a, a, b)]
    assert once == again


def test_a_merge_keeps_the_copy_that_has_the_photograph():
    without = [AssetEvent("2023-05-11", "repair", "new seals")]
    with_photo = [AssetEvent("2023-05-11", "repair", "new seals",
                             photo="data:image/jpeg;base64,AAAA")]
    for order in ((without, with_photo), (with_photo, without)):
        merged = merge_events("X", *order)
        assert len(merged) == 1 and merged[0].photo


# ---------------------------------------------------------------------------
# What is due
# ---------------------------------------------------------------------------

def test_an_annual_sample_falls_due_a_year_after_the_last_one():
    state = asset_state(_asset(
        AssetEvent("2020-01-10", "commissioned"),
        AssetEvent("2023-03-01", "water_sample"),
    ), date(2024, 6, 1))
    sample = [i for i in state.due if i.key == "water_sample"][0]
    assert sample.due_on == date(2024, 3, 1)
    assert sample.state == "overdue"
    assert "2024-03-01" in sample.detail


def test_a_borehole_that_has_never_been_sampled_counts_from_commissioning():
    """Otherwise a point nobody ever sampled is never overdue for sampling."""
    state = asset_state(_asset(AssetEvent("2020-01-10", "commissioned")),
                        date(2024, 6, 1))
    sample = [i for i in state.due if i.key == "water_sample"][0]
    assert sample.state == "overdue"
    assert "has never been done" in sample.detail


def test_with_no_dates_at_all_the_schedule_is_unknown_rather_than_satisfied():
    state = asset_state(_asset(), date(2024, 6, 1))
    assert {i.state for i in state.due} == {"unknown"}
    assert state.overdue == ()


def test_the_grace_period_separates_due_from_overdue():
    events = [AssetEvent("2020-01-10", "commissioned"),
              AssetEvent("2023-06-01", "inspection")]
    due = asset_state(_asset(*events), date(2023, 12, 15))
    overdue = asset_state(_asset(*events), date(2024, 2, 1))
    assert [i.state for i in due.due if i.key == "inspection"] == ["due"]
    assert [i.state for i in overdue.due if i.key == "inspection"] == ["overdue"]


def test_a_due_date_never_drifts_forward_over_a_short_month():
    state = asset_state(_asset(
        AssetEvent("2020-01-10", "commissioned"),
        AssetEvent("2023-08-31", "inspection"),
    ), date(2023, 9, 1))
    inspection = [i for i in state.due if i.key == "inspection"][0]
    assert inspection.due_on == date(2024, 2, 29)


def test_the_programme_sets_its_own_intervals():
    events = [AssetEvent("2020-01-10", "commissioned"),
              AssetEvent("2023-06-01", "water_sample")]
    quarterly = Schedule(water_sample_months=3, grace_days=0)
    state = asset_state(_asset(*events), date(2023, 10, 1), quarterly)
    sample = [i for i in state.due if i.key == "water_sample"][0]
    assert sample.due_on == date(2023, 9, 1) and sample.state == "overdue"


def test_a_decommissioned_borehole_is_not_chased_for_inspections():
    state = asset_state(_asset(
        AssetEvent("2020-01-10", "commissioned"),
        AssetEvent("2022-08-01", "decommissioned", "collapsed, backfilled"),
    ), date(2024, 6, 1))
    assert state.function == "decommissioned"
    assert state.due == ()


# ---------------------------------------------------------------------------
# Many of them
# ---------------------------------------------------------------------------

def test_a_functionality_rate_is_never_computed_over_silence():
    """The number that makes these registers untrustworthy."""
    silent = [_asset() for _ in range(9)]
    working = _asset(AssetEvent("2023-01-01", "commissioned"))
    stats = registry_stats(silent + [working], date(2024, 6, 1))
    assert stats["n_assets"] == 10
    assert stats["n_unknown"] == 9
    assert stats["functionality_rate"] == 100.0   # over the one that is known
    assert stats["n_functional"] == 1

    assert registry_stats([], date(2024, 6, 1))["functionality_rate"] is None


def test_the_headline_counts_add_up():
    assets = [
        _asset(AssetEvent("2023-01-01", "commissioned")),
        _asset(AssetEvent("2023-01-01", "commissioned"),
               AssetEvent("2023-09-01", "failure")),
        _asset(AssetEvent("2023-01-01", "commissioned"),
               AssetEvent("2023-09-01", "decommissioned")),
        _asset(),
    ]
    stats = registry_stats(assets, date(2024, 6, 1))
    assert (stats["n_functional"], stats["n_non_functional"],
            stats["n_decommissioned"], stats["n_unknown"]) == (1, 1, 1, 1)
    assert stats["functionality_rate"] == 50.0
    assert stats["mean_days_out_of_service"] == 274


def test_the_registry_record_is_serialisable():
    import json

    import yaml

    asset = _asset(AssetEvent("2023-05-11", "repair", "new seals",
                              by="A. Bangura"))
    payload = {"asset": asset.as_dict(),
               "state": asset_state(asset, date(2024, 6, 1)).as_dict()}
    assert json.dumps(payload)
    assert yaml.safe_dump(payload)
    restored = registry.asset_from_dict(payload["asset"])
    assert restored is not None
    assert restored.asset_id == asset.asset_id
    assert [e.as_dict() for e in restored.events] == \
        [e.identified_for(asset.asset_id).as_dict() for e in asset.events]


def test_a_saved_record_with_a_broken_identifier_is_not_loaded():
    """It would attach somebody else's history to this borehole."""
    assert registry.asset_from_dict({"asset_id": "SL-WAR-XXXXXXX-9"}) is None
    assert registry.asset_from_dict({"community": "no id at all"}) is None
    assert registry.asset_from_dict("not a record") is None


# ---------------------------------------------------------------------------
# What goes on the borehole
# ---------------------------------------------------------------------------

def test_the_symbol_carries_the_facts_and_not_a_link():
    """A link needs a network. The whole reason it is here is that there isn't one."""
    from groundwater import qr

    asset = _asset()
    asset.total_depth_m, asset.safe_yield_m3_per_h = 62.0, 1.85
    payload = registry.qr_payload(asset)
    assert "http" not in payload
    assert asset.asset_id in payload
    assert "Dr. Timbo's" in payload
    assert "8.48310 N, 13.22940 W" in payload
    assert "62.0 m deep" in payload and "1.85 m3/h" in payload
    qr.encode(payload, ecc="H")           # and it fits in a printable symbol


def test_the_placard_shows_the_condition_when_it_is_known():
    asset = _asset(AssetEvent("2020-01-10", "commissioned"))
    rows = dict(registry.placard_lines(asset, asset_state(asset, date(2024, 6, 1))))
    assert rows["Identifier"] == asset.asset_id
    assert rows["Status"] == "Working"
    assert rows["Commissioned"] == "2020-01-10"


def test_an_asset_is_drafted_from_a_project_but_never_commissioned_by_one():
    """Somebody commissions a borehole. Opening a page does not."""
    asset = registry.asset_from_project({"site": TIMBO})
    assert asset is not None and asset.asset_id == mint_asset_id(TIMBO)
    assert asset.events == []
    assert asset_state(asset, date(2024, 6, 1)).function == "unknown"

    assert registry.asset_from_project({}) is None
    assert registry.asset_from_project(
        {"site": SiteMetadata(community="Nowhere")}) is None


# ---------------------------------------------------------------------------
# The documents
# ---------------------------------------------------------------------------

def test_the_placard_carries_a_symbol_that_reads_back(tmp_path):
    """The plate on the headworks is the whole delivery mechanism."""
    cv2 = pytest.importorskip("cv2", reason="no decoder installed")
    import numpy as np

    from groundwater.reporting.registry import AssetReportInputs, build_asset_placard

    asset = _asset(AssetEvent("2020-01-10", "commissioned"))
    asset.total_depth_m, asset.pump_type = 62.0, "India Mark II"
    out = build_asset_placard(
        AssetReportInputs(asset=asset, figures_dir=tmp_path,
                          today=date(2024, 6, 1)),
        tmp_path / "placard.docx")
    assert out.is_file()

    png = tmp_path / f"qr_{asset.asset_id}.png"
    image = cv2.imdecode(np.frombuffer(png.read_bytes(), dtype=np.uint8),
                         cv2.IMREAD_GRAYSCALE)
    text, *_ = cv2.QRCodeDetector().detectAndDecode(image)
    assert text == registry.qr_payload(asset)
    assert asset.asset_id in text


def test_the_record_says_when_nothing_is_known(tmp_path):
    from docx import Document

    from groundwater.reporting.registry import AssetReportInputs, build_asset_record

    out = build_asset_record(
        AssetReportInputs(asset=_asset(), figures_dir=tmp_path,
                          today=date(2024, 6, 1)),
        tmp_path / "record.docx")
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Not known" in text
    assert "not the same as nothing having happened" in text


def test_the_record_counts_the_days_the_community_went_without(tmp_path):
    from docx import Document

    from groundwater.reporting.registry import AssetReportInputs, build_asset_record

    asset = _asset(AssetEvent("2020-01-10", "commissioned"),
                   AssetEvent("2023-04-02", "failure", "rising main parted"))
    out = build_asset_record(
        AssetReportInputs(asset=asset, figures_dir=tmp_path,
                          today=date(2024, 6, 1)),
        tmp_path / "record.docx")
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Not working" in text
    assert "426 days" in text
    assert "has never been done" in text          # and what is overdue with it
    tables = Document(str(out)).tables
    body = "\n".join(cell.text for t in tables for row in t.rows
                     for cell in row.cells)
    assert "rising main parted" in body


def test_an_asset_record_is_gated_like_every_other_report(tmp_path):
    """A record of a borehole nobody located is not a certification either."""
    from docx import Document

    from groundwater.readiness import assess_readiness
    from groundwater.reporting.registry import AssetReportInputs, build_asset_record

    readiness = assess_readiness({}, "completion")
    out = build_asset_record(
        AssetReportInputs(asset=_asset(), figures_dir=tmp_path,
                          today=date(2024, 6, 1), readiness=readiness),
        tmp_path / "record.docx")
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "PROVISIONAL - NOT FOR CERTIFICATION" in text


def test_a_minted_identifier_is_always_accepted_by_the_parser():
    """Round trip every identifier a stretch of the country mints.

    The check character comes from the full 36-character alphabet, in which
    I, L, O and U mean themselves. Folding them the way the position code
    folds them - where the alphabet has no such letters, so folding is
    unambiguous - rejected one identifier in nine that this same module had
    just minted, and with it the whole maintenance history filed under it.
    """
    from groundwater.models import SiteMetadata
    from groundwater.registry import mint_asset_id, parse_asset_id

    minted = 0
    for easting in range(700000, 700400, 7):
        for northing in range(900000, 900400, 11):
            site = SiteMetadata(district="Port Loko", easting=easting,
                                northing=northing, utm_zone=28)
            asset_id = mint_asset_id(site)
            minted += 1
            assert parse_asset_id(asset_id) == asset_id, asset_id
    assert minted > 2000


def test_a_mistyped_check_character_is_still_refused():
    """Forgiving the four Crockford confusions must not forgive a wrong code."""
    from groundwater.registry import parse_asset_id

    # SL-PL-8HE2QWG-3 is well formed; every other check character is not
    assert parse_asset_id("SL-PL-8HE2QWG-3") == "SL-PL-8HE2QWG-3"
    for wrong in "012456789ABCDEFGHIJKLMNOPQRSTUVWXYZ":
        assert parse_asset_id(f"SL-PL-8HE2QWG-{wrong}") is None


def test_an_emoji_in_a_note_hashes_the_same_as_the_browser():
    """The browser hashes UTF-16 code units; Python used to hash code points.

    They agree on everything in the Basic Multilingual Plane and diverged on
    an emoji, so the same repair noted with a thumbs-up on a phone and typed
    in the office never merged. The literals are the browser engine's own
    answers (docs/js/gwt-core.js eventId)."""
    from groundwater.registry import event_id

    asset = "SL-WAR-8T4KQ2A-7"
    assert event_id(asset, "2024-01-01", "failure", "pump seized", "M") == "2024-01-01/failure/8Z6E"
    assert event_id(asset, "2024-01-01", "failure", "café", "M") == "2024-01-01/failure/GTTD"
    assert event_id(asset, "2024-01-01", "failure", "pump seized 😀", "M") == "2024-01-01/failure/1QW1"
