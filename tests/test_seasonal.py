"""The seasonal yield model.

A pumping test measures one day; the borehole has to supply the village on
the worst day, and in Sierra Leone those are months apart. The point of
these tests is that the month of the test changes the answer, that an
ambiguous date is treated as no date rather than as a guess, and that the
pump is set for the worst scenario rather than the average one.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from groundwater.config import PumpingConfig
from groundwater.hydraulics import analyse_pumping_test
from groundwater.ingestion import read_pumping_workbook
from groundwater.seasonal import (
    SEASONAL_POSITION,
    month_of,
    season_of,
    seasonal_yield,
)


@pytest.fixture(scope="module")
def analysis(sample_data):
    path = Path(sample_data) / "dr_timbo" / "dr_timbo_constant_test.xlsx"
    return analyse_pumping_test(read_pumping_workbook(path))


# ---------------------------------------------------------------------------
# Reading the date
# ---------------------------------------------------------------------------

def test_a_date_that_could_be_read_either_way_is_not_guessed_at():
    """10/05/2018 is 10 May or 5 October - opposite ends of the year."""
    month, why = month_of("10/05/2018")
    assert month is None
    assert "either way round" in why
    assert "May" in why and "October" in why


def test_a_date_with_only_one_reading_is_read():
    assert month_of("25/04/2018")[0] == 4       # 25 cannot be a month
    assert month_of("04/25/2018")[0] == 4       # nor here
    assert month_of("2018-09-14")[0] == 9       # ISO is unambiguous


def test_a_month_written_by_name_settles_it():
    for text in ("14 Sept 2018", "September 2018", "2018 sep 14", "SEPTEMBER"):
        assert month_of(text)[0] == 9, text


def test_no_date_is_reported_as_no_date():
    month, why = month_of("")
    assert month is None and "No date" in why
    month, why = month_of("during the rains")
    assert month is None and "does not name a month" in why
    month, why = month_of("05/2018")
    assert month is None and "not a full date" in why


def test_the_year_has_one_high_and_one_low():
    """A single wet season, not two - this is not equatorial East Africa."""
    assert set(SEASONAL_POSITION) == set(range(1, 13))
    assert all(0.0 <= v <= 1.0 for v in SEASONAL_POSITION.values())
    assert min(SEASONAL_POSITION, key=SEASONAL_POSITION.get) == 9   # recovered
    assert max(SEASONAL_POSITION, key=SEASONAL_POSITION.get) == 5   # driest
    assert season_of(9) == "wet season" and season_of(4) == "late dry season"
    assert season_of(None) == "unknown"


# ---------------------------------------------------------------------------
# What the scenarios say
# ---------------------------------------------------------------------------

def test_a_wet_season_test_loses_yield_by_the_end_of_the_dry(analysis):
    """The whole point: August's number is not the number that matters."""
    result = seasonal_yield(analysis, month=8)
    assert result.is_established
    tested = result.scenario("as_tested").safe_yield_m3_per_h
    design = result.scenario("dry_season").safe_yield_m3_per_h
    drought = result.scenario("drought").safe_yield_m3_per_h
    assert tested > design > drought
    assert result.dry_season_loss_percent > 0
    assert result.design_yield_m3_per_h == design


def test_a_test_run_at_the_annual_low_reserves_nothing_further(analysis):
    """A May test has already spent the decline; taking it twice is wrong."""
    result = seasonal_yield(analysis, month=5)
    assert result.scenario("dry_season").decline_m == 0.0
    assert (result.scenario("as_tested").safe_yield_m3_per_h
            == result.scenario("dry_season").safe_yield_m3_per_h)
    assert "already at or below this level" in result.scenario("dry_season").note
    # the drought case still reserves something: a bad year goes lower
    assert result.scenario("drought").decline_m > 0


def test_the_same_test_in_may_beats_the_same_test_in_september(analysis):
    """It is the same borehole; a dry-season test simply proves more."""
    may = seasonal_yield(analysis, month=5).design_yield_m3_per_h
    september = seasonal_yield(analysis, month=9).design_yield_m3_per_h
    assert may > september


def test_an_unknown_month_reserves_the_whole_range(analysis):
    """The conservative reading, said out loud rather than assumed."""
    unknown = seasonal_yield(analysis, month=None, annual_range_m=2.0)
    september = seasonal_yield(analysis, month=9, annual_range_m=2.0)
    assert unknown.season == "unknown"
    assert (unknown.scenario("dry_season").decline_m
            == september.scenario("dry_season").decline_m == 2.0)


def test_the_field_sheet_date_is_used_when_it_can_be(analysis):
    """And its ambiguity is carried into the result rather than resolved."""
    result = seasonal_yield(analysis)
    date = analysis.test.site.date
    expected, why = month_of(date)
    assert result.month == expected
    assert result.month_note == why
    if expected is None:
        assert result.season == "unknown"


def test_the_pump_is_set_for_the_worst_year_not_the_average_one(analysis):
    """It is fitted once. A pump that is dry in a drought is a lost borehole."""
    result = seasonal_yield(analysis, month=9)
    depths = {s.key: s.pump_installation_depth_m for s in result.scenarios}
    assert depths["drought"] >= depths["dry_season"] >= depths["as_tested"]
    assert result.pump_installation_depth_m == depths["drought"]


def test_the_range_is_a_parameter_and_says_where_it_came_from(analysis):
    assumed = seasonal_yield(analysis, month=9)
    assert "not a measurement" in assumed.range_source
    assert assumed.annual_range_m == PumpingConfig().seasonal_allowance_m

    measured = seasonal_yield(analysis, month=9, annual_range_m=4.5)
    assert measured.annual_range_m == 4.5
    assert "measured" in measured.range_source
    # a bigger swing takes more away
    assert measured.design_yield_m3_per_h < assumed.design_yield_m3_per_h


def test_a_pending_recommendation_stays_pending(analysis):
    """Nothing is invented from a test that could not be analysed."""
    class Empty:
        test = None
        yield_recommendation = None

    result = seasonal_yield(Empty())
    assert not result.is_established
    assert result.design_yield_m3_per_h is None
    assert "nothing to project" in result.pending_reason
    assert result.summary == result.pending_reason
    assert result.scenarios == []


def test_the_summary_names_the_month_and_both_figures(analysis):
    result = seasonal_yield(analysis, month=8)
    assert "end of the dry season" in result.summary
    assert "drought year" in result.summary
    assert "August" in result.summary and "wet season" in result.summary


def test_the_seasonal_record_is_serialisable(analysis):
    import json

    import yaml

    payload = seasonal_yield(analysis, month=8).as_dict()
    assert json.dumps(payload)
    assert yaml.safe_dump(payload)
    assert len(payload["scenarios"]) == 3


# ---------------------------------------------------------------------------
# What the report says
# ---------------------------------------------------------------------------

def test_the_pumping_report_shows_the_yield_through_the_year(analysis, tmp_path):
    from docx import Document

    from groundwater.reporting.pumping import (
        PumpingReportInputs,
        build_pumping_report,
    )

    result = seasonal_yield(analysis, month=8)
    out = build_pumping_report(
        PumpingReportInputs(analysis=analysis, figures_dir=tmp_path,
                            seasonal=result),
        tmp_path / "pumping.docx")
    document = Document(str(out))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Through the year" in text
    assert "August" in text and "wet season" in text
    assert "the pump is fitted once" in text
    assert f"{result.annual_range_m:.1f} m" in text

    body = "\n".join(cell.text for t in document.tables for row in t.rows
                     for cell in row.cells)
    for scenario in result.scenarios:
        assert scenario.title in body


def test_an_ambiguous_date_is_explained_in_the_report(analysis, tmp_path):
    """Rather than the report quietly picking one end of the year."""
    from docx import Document

    from groundwater.reporting.pumping import (
        PumpingReportInputs,
        build_pumping_report,
    )

    out = build_pumping_report(
        PumpingReportInputs(analysis=analysis, figures_dir=tmp_path,
                            seasonal=seasonal_yield(analysis, month=None)),
        tmp_path / "pumping_unknown.docx")
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "conservative reading" in text


def test_a_report_without_the_projection_is_unchanged(analysis, tmp_path):
    """An existing caller that passes nothing gets the report it always got."""
    from docx import Document

    from groundwater.reporting.pumping import (
        PumpingReportInputs,
        build_pumping_report,
    )

    out = build_pumping_report(
        PumpingReportInputs(analysis=analysis, figures_dir=tmp_path),
        tmp_path / "pumping_plain.docx")
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Through the year" not in text
